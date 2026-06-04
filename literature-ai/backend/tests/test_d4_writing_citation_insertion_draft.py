from __future__ import annotations

import tempfile
from io import BytesIO
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, func, select
from sqlalchemy.orm import sessionmaker

from app.config import get_settings
from app.db.models import (
    Base,
    DFTResult,
    EvidenceLocator,
    ExtractionFieldReview,
    Paper,
    PaperCitationEligibility,
    PaperImpactMetadata,
)
from app.db.session import get_db_session
from app.main import app


TEXT = "Single-atom catalysts can accelerate sulfur redox kinetics in lithium-sulfur batteries."


@pytest.fixture
def insertion_client(monkeypatch):
    with tempfile.TemporaryDirectory() as tmpdir:
        db_path = Path(tmpdir) / "citation_insertion.db"
        db_url = f"sqlite:///{db_path}"
        monkeypatch.setenv("LITAI_DATABASE_URL", db_url)
        monkeypatch.setenv("LITAI_STORAGE_ROOT", str(Path(tmpdir) / "storage"))
        get_settings.cache_clear()
        engine = create_engine(db_url, future=True)
        Base.metadata.create_all(engine)
        Session = sessionmaker(autocommit=False, autoflush=False, bind=engine, future=True)

        def override_get_db_session():
            db = Session()
            try:
                yield db
            finally:
                db.close()

        app.dependency_overrides[get_db_session] = override_get_db_session
        seed = _seed(Session)
        yield TestClient(app), Session, seed
        app.dependency_overrides.clear()
        engine.dispose()
        from app.db.session import _engines, _session_factories

        for cached_engine in list(_engines.values()):
            cached_engine.dispose()
        _engines.clear()
        _session_factories.clear()
        get_settings.cache_clear()


def test_blank_text_returns_422(insertion_client):
    client, _, seed = insertion_client
    response = _post(client, seed["metadata_only"], text="   ")
    assert response.status_code == 422


def test_missing_paper_returns_404(insertion_client):
    client, _, _ = insertion_client
    response = _post(client, uuid4())
    assert response.status_code == 404
    assert response.json()["detail"] == "Paper not found"


def test_metadata_only_returns_draft_with_human_verification(insertion_client):
    client, _, seed = insertion_client
    response = _post(client, seed["metadata_only"])
    assert response.status_code == 200
    data = response.json()
    assert data["proposal_status"] == "metadata_only_draft"
    assert data["can_insert_as_confirmed_citation"] is False
    assert data["requires_human_verification"] is True
    assert any("Metadata-only suggestion cannot be used as evidence yet" in warning for warning in data["warnings"])
    assert "[DRAFT CITATION - VERIFY SOURCE BEFORE USE:" in data["draft_text"]


def test_unverified_extraction_requires_verification_marker(insertion_client):
    client, _, seed = insertion_client
    response = _post(client, seed["unverified_extraction"])
    assert response.status_code == 200
    data = response.json()
    assert data["proposal_status"] == "needs_human_verification"
    assert data["evidence_status"] == "unverified_extraction"
    assert data["can_insert_as_confirmed_citation"] is False
    assert "VERIFY SOURCE BEFORE USE" in data["draft_text"]
    assert any("VERIFY SOURCE BEFORE USE" in warning for warning in data["warnings"])


def test_safe_verified_can_be_confirmed_only_when_db_safe(insertion_client):
    client, _, seed = insertion_client
    response = _post(client, seed["safe"])
    assert response.status_code == 200
    data = response.json()
    assert data["proposal_status"] == "confirmed_candidate_draft"
    assert data["evidence_status"] == "safe_verified"
    assert data["can_insert_as_confirmed_citation"] is True
    assert data["requires_human_verification"] is False
    assert data["draft_text"].endswith("(Safe et al., 2024).")


def test_verified_but_not_safe_requires_safety_review(insertion_client):
    client, _, seed = insertion_client
    response = _post(client, seed["verified_unsafe"])
    assert response.status_code == 200
    data = response.json()
    assert data["proposal_status"] == "verified_but_requires_safety_review"
    assert data["evidence_status"] == "verified"
    assert data["can_insert_as_confirmed_citation"] is False
    assert any("Verified does not equal safe_verified" in warning for warning in data["warnings"])


def test_forged_frontend_confirmed_flag_is_ignored(insertion_client):
    client, _, seed = insertion_client
    response = _post(
        client,
        seed["metadata_only"],
        candidate_can_be_used_as_confirmed_citation=True,
        candidate_evidence_status="safe_verified",
    )
    assert response.status_code == 200
    data = response.json()
    assert data["evidence_status"] == "metadata_only"
    assert data["can_insert_as_confirmed_citation"] is False
    assert any("Client-provided confirmed citation flag was ignored" in warning for warning in data["warnings"])


def test_excluded_paper_is_hard_blocked(insertion_client):
    client, _, seed = insertion_client
    response = _post(client, seed["excluded"])
    assert response.status_code == 200
    data = response.json()
    assert data["proposal_status"] == "blocked_excluded_from_citation"
    assert data["blocked_reason"] == "exclude_from_citation=true"
    assert data["draft_text"] is None
    assert data["can_insert_as_confirmed_citation"] is False


def test_citation_priority_exclude_is_hard_blocked(insertion_client):
    client, _, seed = insertion_client
    response = _post(client, seed["priority_exclude"])
    assert response.status_code == 200
    data = response.json()
    assert data["proposal_status"] == "blocked_excluded_from_citation"
    assert data["blocked_reason"] == "citation_priority=exclude"
    assert data["draft_text"] is None


def test_api_does_not_write_database(insertion_client):
    client, Session, seed = insertion_client
    with Session() as session:
        before = _counts(session)
        row_before = _row_snapshot(session, seed["safe"])
    response = _post(client, seed["safe"])
    assert response.status_code == 200
    with Session() as session:
        assert _counts(session) == before
        assert _row_snapshot(session, seed["safe"]) == row_before


def test_response_does_not_include_bibliography(insertion_client):
    client, _, seed = insertion_client
    response = _post(client, seed["safe"])
    assert response.status_code == 200
    assert "bibliography" not in response.json()
    assert "references" not in response.json()


def test_word_insert_appends_safe_verified_citation_copy(insertion_client):
    client, Session, seed = insertion_client
    with Session() as session:
        before = _counts(session)
    response = _post_word(
        client,
        seed["safe"],
        document_bytes=_docx_bytes(["Draft manuscript body."]),
        output_filename="safe-citation-output.docx",
    )
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "inserted"
    assert data["output_filename"] == "safe-citation-output.docx"
    assert data["output_relative_path"] == "word_exports/safe-citation-output.docx"
    assert data["download_url"] == "/api/writing/word/exports/safe-citation-output.docx"
    assert data["safety"]["mutates_original_file"] is False
    assert data["safety"]["writes_database"] is False
    assert data["safety"]["can_insert_as_confirmed_citation"] is True
    paragraphs = _docx_paragraphs(Path(data["output_path"]))
    assert paragraphs[0] == "Draft manuscript body."
    assert paragraphs[-1].endswith("(Safe et al., 2024).")
    downloaded = client.get(data["download_url"])
    assert downloaded.status_code == 200
    assert downloaded.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert downloaded.content.startswith(b"PK")
    with Session() as session:
        assert _counts(session) == before


def test_word_insert_replaces_placeholder_with_unverified_marker(insertion_client):
    client, _, seed = insertion_client
    response = _post_word(
        client,
        seed["metadata_only"],
        document_bytes=_docx_bytes([f"{TEXT} {{CITE}}"]),
        docx_insertion_mode="replace_placeholder",
        placeholder="{CITE}",
        output_filename="placeholder-output.docx",
    )
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "inserted"
    assert data["placeholder_replaced_count"] == 1
    assert data["safety"]["requires_human_verification"] is True
    assert "[DRAFT CITATION - VERIFY SOURCE BEFORE USE:" in data["inserted_text"]
    paragraphs = _docx_paragraphs(Path(data["output_path"]))
    assert "{CITE}" not in paragraphs[0]
    assert "[DRAFT CITATION - VERIFY SOURCE BEFORE USE:" in paragraphs[0]


def test_word_insert_excluded_paper_is_blocked_without_output(insertion_client):
    client, _, seed = insertion_client
    response = _post_word(client, seed["excluded"], document_bytes=_docx_bytes(["Draft manuscript body."]))
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "blocked"
    assert data["output_path"] is None
    assert data["download_url"] is None
    assert data["draft"]["proposal_status"] == "blocked_excluded_from_citation"
    assert data["inserted_text"] is None


def _post(
    client,
    paper_id,
    *,
    text=TEXT,
    candidate_can_be_used_as_confirmed_citation=False,
    candidate_evidence_status="metadata_only",
):
    return client.post(
        "/api/writing/citation-insertion-draft",
        json={
            "text": text,
            "selected_paper_id": str(paper_id),
            "citation_marker": "",
            "insertion_mode": "parenthetical",
            "citation_style": "draft_author_year",
            "candidate_evidence_status": candidate_evidence_status,
            "candidate_can_be_used_as_confirmed_citation": candidate_can_be_used_as_confirmed_citation,
            "candidate_requires_human_verification": True,
            "supporting_snippet": "client supplied context",
            "user_note": "",
        },
    )


def _post_word(
    client,
    paper_id,
    *,
    document_bytes,
    text=TEXT,
    docx_insertion_mode="append_paragraph",
    citation_insertion_mode="parenthetical",
    placeholder=None,
    output_filename="citation-output.docx",
):
    data = {
        "text": text,
        "selected_paper_id": str(paper_id),
        "docx_insertion_mode": docx_insertion_mode,
        "citation_insertion_mode": citation_insertion_mode,
        "citation_style": "draft_author_year",
        "output_filename": output_filename,
    }
    if placeholder is not None:
        data["placeholder"] = placeholder
    return client.post(
        "/api/writing/word/insert-citation",
        files={
            "file": (
                "draft.docx",
                document_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data=data,
    )


def _docx_bytes(paragraphs):
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _docx_paragraphs(path: Path):
    return [paragraph.text for paragraph in Document(path).paragraphs]


def _seed(Session):
    with Session() as session:
        metadata_only = Paper(
            title="Single-atom sulfur redox kinetics overview",
            year=2024,
            journal="Metadata Journal",
            abstract="Single-atom catalysts and sulfur redox kinetics in lithium-sulfur batteries.",
            authors=[{"last": "Meta"}],
            pdf_path="meta.pdf",
        )
        unverified = Paper(
            title="Unverified extraction",
            year=2024,
            journal="Extraction Journal",
            abstract="battery catalysis",
            authors=[{"last": "Extract"}],
            pdf_path="extract.pdf",
        )
        safe = Paper(
            title="Safe verified",
            year=2024,
            journal="Safe Journal",
            abstract="battery catalysis",
            authors=[{"last": "Safe"}, {"last": "Coauthor"}],
            pdf_path="safe.pdf",
        )
        verified_unsafe = Paper(
            title="Verified unsafe",
            year=2024,
            journal="Unsafe Journal",
            abstract="battery catalysis",
            authors=[{"last": "Unsafe"}],
            pdf_path="unsafe.pdf",
        )
        excluded = Paper(
            title="Do not cite",
            year=2024,
            journal="Excluded Journal",
            abstract="single-atom sulfur redox kinetics",
            authors=[{"last": "Excluded"}],
            pdf_path="excluded.pdf",
        )
        priority_exclude = Paper(
            title="Priority exclude",
            year=2024,
            journal="Excluded Journal",
            abstract="single-atom sulfur redox kinetics",
            authors=[{"last": "Priority"}],
            pdf_path="priority.pdf",
        )
        session.add_all([metadata_only, unverified, safe, verified_unsafe, excluded, priority_exclude])
        session.flush()
        session.add_all(
            [
                PaperImpactMetadata(paper_id=unverified.id, impact_factor=5.0, impact_factor_source="fixture", impact_factor_year=2024),
                PaperImpactMetadata(paper_id=safe.id, impact_factor=12.0, impact_factor_source="fixture", impact_factor_year=2024),
                PaperImpactMetadata(paper_id=verified_unsafe.id, impact_factor=13.0, impact_factor_source="fixture", impact_factor_year=2024),
                PaperCitationEligibility(paper_id=safe.id, citation_priority="high"),
                PaperCitationEligibility(paper_id=excluded.id, exclude_from_citation=True, exclude_reason="manual"),
                PaperCitationEligibility(paper_id=priority_exclude.id, citation_priority="exclude"),
            ]
        )
        unverified_result = _result(unverified.id)
        safe_result = _result(safe.id)
        unsafe_result = _result(verified_unsafe.id)
        session.add_all([unverified_result, safe_result, unsafe_result])
        session.flush()
        session.add_all(
            [
                _review(safe.id, safe_result.id, "verified", "active"),
                _review(verified_unsafe.id, unsafe_result.id, "verified", "stale"),
                _locator(safe.id, safe_result.id, "exact", 4),
                _locator(verified_unsafe.id, unsafe_result.id, "exact", 8),
            ]
        )
        session.commit()
        return {
            "metadata_only": metadata_only.id,
            "unverified_extraction": unverified.id,
            "safe": safe.id,
            "verified_unsafe": verified_unsafe.id,
            "excluded": excluded.id,
            "priority_exclude": priority_exclude.id,
        }


def _result(paper_id):
    return DFTResult(
        paper_id=paper_id,
        property_type="sulfur redox kinetics",
        value=1.0,
        unit="a.u.",
        evidence_text="Single-atom catalysts accelerate sulfur redox kinetics in lithium-sulfur batteries.",
    )


def _review(paper_id, target_id, reviewer_status, target_resolution_status):
    return ExtractionFieldReview(
        paper_id=paper_id,
        target_type="dft_results",
        target_id=str(target_id),
        field_name="value",
        original_value=1.0,
        reviewed_value=1.0,
        evidence_text="Single-atom catalysts accelerate sulfur redox kinetics in lithium-sulfur batteries.",
        reviewer_status=reviewer_status,
        target_resolution_status=target_resolution_status,
    )


def _locator(paper_id, target_id, locator_status, page):
    return EvidenceLocator(
        paper_id=paper_id,
        target_type="dft_results",
        target_id=str(target_id),
        source_type="text",
        page=page,
        evidence_text="Single-atom catalysts accelerate sulfur redox kinetics in lithium-sulfur batteries.",
        locator_status=locator_status,
        locator_confidence=0.8,
        parser_source="fixture",
    )


def _counts(session):
    return {
        "papers": session.scalar(select(func.count(Paper.id))),
        "reviews": session.scalar(select(func.count(ExtractionFieldReview.id))),
        "locators": session.scalar(select(func.count(EvidenceLocator.id))),
        "eligibility": session.scalar(select(func.count(PaperCitationEligibility.paper_id))),
        "impact": session.scalar(select(func.count(PaperImpactMetadata.paper_id))),
    }


def _row_snapshot(session, paper_id):
    paper = session.get(Paper, paper_id)
    review = session.scalar(select(ExtractionFieldReview).where(ExtractionFieldReview.paper_id == paper_id))
    locator = session.scalar(select(EvidenceLocator).where(EvidenceLocator.paper_id == paper_id))
    eligibility = session.get(PaperCitationEligibility, paper_id)
    impact = session.get(PaperImpactMetadata, paper_id)
    return {
        "paper": (paper.title, paper.year, paper.journal, paper.abstract, paper.pdf_path),
        "review": (review.reviewer_status, review.target_resolution_status, review.evidence_text),
        "locator": (locator.locator_status, locator.page, locator.evidence_text),
        "eligibility": (eligibility.exclude_from_citation, eligibility.citation_priority),
        "impact": (impact.impact_factor, impact.impact_factor_year, impact.impact_factor_source),
    }
