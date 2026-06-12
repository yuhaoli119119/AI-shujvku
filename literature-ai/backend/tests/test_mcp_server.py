from __future__ import annotations

import tempfile
from pathlib import Path
from uuid import UUID

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session, sessionmaker

from app.config import get_settings
from app.db.models import (
    AuditLog,
    Base,
    CatalystSample,
    DFTResult,
    ElectrochemicalPerformance,
    EvidenceLocator,
    ExtractionFieldReview,
    ExternalAnalysisCandidate,
    ExternalAnalysisRun,
    MechanismClaim,
    Paper,
    PaperCorrection,
    PaperFigure,
    PaperNote,
    PaperSection,
    PaperTable,
    WritingCard,
)
from app.main import app
from app.mcp.auth import parse_mcp_api_keys
from app.mcp.context import MCPAuthInfo, mcp_auth_context
from app.mcp.server import (
    append_note,
    approve_correction,
    get_correction_detail,
    get_codex_context,
    get_correction_queue,
    get_codex_item,
    get_dft_review_queue,
    get_paper_knowledge,
    get_parse_status,
    ingest_pdf_batch,
    insert_word_citation,
    import_analysis,
    list_notes,
    parse_paper,
    propose_correction,
    propose_dft_result_correction,
    reject_dft_result,
    query_papers,
    reject_correction,
    scan_local_pdfs,
    scan_duplicate_dois,
)
from app.utils.library_names import DEFAULT_LIBRARY_NAME


@pytest.fixture
def mcp_test_env(monkeypatch):
    with tempfile.TemporaryDirectory() as tmpdir:
        db_path = Path(tmpdir) / "mcp_test.db"
        db_url = f"sqlite:///{db_path}"
        monkeypatch.setenv("LITAI_DATABASE_URL", db_url)
        monkeypatch.setenv(
            "LITAI_MCP_API_KEYS",
            "claude|Claude Desktop|litmcp_claude|read_papers,append_notes,propose_corrections,request_parse;"
            "admin|Admin|litmcp_admin|read_papers,append_notes,propose_corrections,request_parse,review_corrections",
        )
        monkeypatch.setenv("LITAI_STORAGE_ROOT", str(Path(tmpdir) / "storage"))
        get_settings.cache_clear()

        engine = create_engine(db_url, future=True)
        Base.metadata.create_all(engine)
        SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False, future=True)

        yield {"sessionmaker": SessionLocal, "engine": engine, "tmpdir": Path(tmpdir)}

        engine.dispose()
        from app.db.session import _engines, _session_factories

        for eng in list(_engines.values()):
            try:
                eng.dispose()
            except Exception:
                pass
        _engines.clear()
        _session_factories.clear()
        get_settings.cache_clear()


def _auth() -> MCPAuthInfo:
    return MCPAuthInfo(
        source_prefix="claude",
        display_name="Claude Desktop",
        capabilities=frozenset({"read_papers", "append_notes", "propose_corrections", "request_parse"}),
        raw_key="litmcp_claude",
    )


def _ide_auth() -> MCPAuthInfo:
    return MCPAuthInfo(
        source_prefix="ide_ai",
        display_name="IDE AI",
        capabilities=frozenset({"read_papers", "append_notes", "propose_corrections", "request_parse"}),
        raw_key="litmcp_ide_ai",
    )


def _admin_auth() -> MCPAuthInfo:
    return MCPAuthInfo(
        source_prefix="admin",
        display_name="Admin",
        capabilities=frozenset(
            {"read_papers", "append_notes", "propose_corrections", "request_parse", "review_corrections"}
        ),
        raw_key="litmcp_admin",
    )


def _make_external_audit_ready(paper: Paper, root: Path) -> None:
    pdf_path = root / f"{paper.id}.pdf"
    markdown_path = root / f"{paper.id}.md"
    docling_path = root / f"{paper.id}.docling.json"
    workspace_path = root / "workspace" / str(paper.id)
    pdf_path.write_bytes(b"%PDF-1.4\n1 0 obj<<>>endobj\ntrailer<<>>\n%%EOF\n")
    markdown_path.write_text("# Ready paper\n\nDFT evidence is available.", encoding="utf-8")
    docling_path.write_text('{"texts": [{"text": "DFT evidence is available."}]}', encoding="utf-8")
    package_path = workspace_path / "extraction" / "ai_reading_package.json"
    package_path.parent.mkdir(parents=True, exist_ok=True)
    package_path.write_text('{"sections": [{"title": "Results"}]}', encoding="utf-8")
    paper.pdf_path = str(pdf_path)
    paper.markdown_path = str(markdown_path)
    paper.docling_json_path = str(docling_path)
    paper.workspace_path = str(workspace_path)


def test_mcp_query_note_and_correction_workflow(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(
            doi="10.1000/test-doi",
            title="MCP Test Paper",
            journal="Nature Energy",
            year=2025,
            authors=["Alice", "Bob"],
            pdf_path="test.pdf",
        )
        session.add(paper)
        session.commit()
        paper_id = str(paper.id)

    with mcp_auth_context(_auth()):
        result = query_papers(q="MCP Test", limit=10)
        assert result["returned"] == 1
        assert result["items"][0]["title"] == "MCP Test Paper"

        note = append_note(
            paper_id=paper_id,
            content="The adsorption energy sentence should be rechecked.",
            field_name="dft_results_items",
            page=5,
            section_title="Results and Discussion",
            quoted_text="The adsorption energy of Li2S4 is -1.23 eV.",
        )
        assert note["source"] == "claude"
        assert note["page"] == 5

        notes = list_notes(paper_id=paper_id)
        assert len(notes["items"]) == 1
        assert notes["items"][0]["quoted_text"] == "The adsorption energy of Li2S4 is -1.23 eV."

        correction = propose_correction(
            paper_id=paper_id,
            field_name="abstract",
            target_path="abstract",
            operation="replace",
            proposed_value="Updated abstract text",
            reason="Cross-check against the uploaded PDF abstract.",
            evidence_payload={"page": 1, "section_title": "Abstract"},
        )
        assert correction["status"] == "pending"
        assert correction["target_path"] == "abstract"

    with Session(mcp_test_env["engine"]) as session:
        saved_notes = session.scalars(select(PaperNote)).all()
        saved_corrections = session.scalars(select(PaperCorrection)).all()
        audit_logs = session.scalars(select(AuditLog).order_by(AuditLog.created_at.asc())).all()

        assert len(saved_notes) == 1
        assert len(saved_corrections) == 1
        assert [item.action for item in audit_logs] == ["append_note", "propose_correction"]


def test_scan_duplicate_dois_groups_default_library_aliases(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        session.add_all(
            [
                Paper(title="Default Alias One", doi="10.1000/default-alias", library_name=DEFAULT_LIBRARY_NAME, pdf_path="one.pdf"),
                Paper(title="Default Alias Two", doi="10.1000/default-alias", library_name="Codex ????????", pdf_path="two.pdf"),
                Paper(title="Other Library", doi="10.1000/default-alias", library_name="OtherLibrary", pdf_path="three.pdf"),
            ]
        )
        session.commit()

    with mcp_auth_context(_auth()):
        payload = scan_duplicate_dois()

    duplicate = next(item for item in payload["duplicates"] if item["doi"] == "10.1000/default-alias")
    assert duplicate["library_name"] == DEFAULT_LIBRARY_NAME
    assert duplicate["count"] == 2
    assert len(duplicate["paper_ids"]) == 2


def test_mcp_query_papers_sort_by_created_at(mcp_test_env):
    """Verify query_papers supports sort_by='created_at' and sort_order."""
    from datetime import datetime, timezone

    with Session(mcp_test_env["engine"]) as session:
        paper1 = Paper(
            doi="10.1000/old",
            title="Old Paper",
            year=2020,
            pdf_path="old.pdf",
        )
        session.add(paper1)
        session.flush()
        # Force older created_at
        paper1.created_at = datetime(2024, 1, 1, 0, 0, 0, tzinfo=timezone.utc)

        paper2 = Paper(
            doi="10.1000/new",
            title="New Paper",
            year=2025,
            pdf_path="new.pdf",
        )
        session.add(paper2)
        session.flush()
        paper2.created_at = datetime(2025, 1, 1, 0, 0, 0, tzinfo=timezone.utc)
        session.commit()

    with mcp_auth_context(_auth()):
        # Descending: newest first
        desc = query_papers(sort_by="created_at", sort_order="desc", limit=10)
        assert desc["returned"] == 2
        assert desc["items"][0]["title"] == "New Paper"
        assert desc["items"][1]["title"] == "Old Paper"

        # Ascending: oldest first
        asc = query_papers(sort_by="created_at", sort_order="asc", limit=10)
        assert asc["returned"] == 2
        assert asc["items"][0]["title"] == "Old Paper"
        assert asc["items"][1]["title"] == "New Paper"


def test_mcp_get_codex_item_returns_low_token_dft_context(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(title="MCP Codex Item Paper", pdf_path="codex-item.pdf")
        session.add(paper)
        session.flush()
        row = DFTResult(
            paper_id=paper.id,
            property_type="formation_energy",
            value=7.5,
            unit="eV",
            evidence_text="The reported defect formation energy is 7.5 eV.",
        )
        session.add(row)
        session.flush()
        session.add(
            EvidenceLocator(
                paper_id=paper.id,
                target_type="dft_result",
                target_id=str(row.id),
                field_name="value",
                page=5,
                evidence_text=row.evidence_text,
                locator_status="exact_page",
                locator_confidence=0.9,
                parser_source="test",
            )
        )
        session.commit()
        paper_id = str(paper.id)
        row_id = str(row.id)

    with mcp_auth_context(_auth()):
        payload = get_codex_item(
            paper_id=paper_id,
            item_type="dft_result",
            item_id=row_id,
        )

    assert payload["schema_version"] == "codex_item_context_v1"
    assert payload["item_type"] == "dft_result"
    assert payload["context"]["item"]["value"] == 7.5
    assert payload["context"]["export_safety"]["blocked_reasons"] == ["missing_review"]
    assert payload["context"]["evidence_locators"]["items"][0]["page"] == 5


def test_ordinary_ide_ai_reads_context_and_imports_unverified_audit_candidate(mcp_test_env):
    configs = parse_mcp_api_keys(
        "ide_ai|IDE AI|litmcp_ide_ai|read_papers,append_notes,propose_corrections,request_parse"
    )
    assert configs["litmcp_ide_ai"].capabilities == frozenset(
        {"read_papers", "append_notes", "propose_corrections", "request_parse"}
    )
    assert "review_corrections" not in configs["litmcp_ide_ai"].capabilities

    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(
            title="Ordinary IDE AI MCP Workflow Paper",
            doi="10.1000/ide-ai-workflow",
            year=2026,
            journal="Workflow Journal",
            authors=["AI Reviewer"],
            abstract="A paper with DFT, figures, tables, mechanism claims, and writing cards.",
            pdf_path="workflow.pdf",
        )
        session.add(paper)
        session.flush()
        _make_external_audit_ready(paper, mcp_test_env["tmpdir"])
        section = PaperSection(
            paper_id=paper.id,
            section_title="Results",
            section_type="results",
            text="The adsorption energy of Li2S4 is -1.23 eV and the figure supports the trend.",
            page_start=3,
            page_end=4,
        )
        figure = PaperFigure(
            paper_id=paper.id,
            caption="Figure 2. Adsorption configuration and charge redistribution.",
            image_path="figures/fig2.png",
            page=4,
            figure_role="data_figure",
        )
        table = PaperTable(
            paper_id=paper.id,
            caption="Table 1. DFT adsorption energies.",
            markdown_content="| Species | Energy |\n| Li2S4 | -1.23 eV |",
            page=5,
        )
        dft = DFTResult(
            paper_id=paper.id,
            property_type="adsorption_energy",
            adsorbate="Li2S4",
            value=-1.23,
            unit="eV",
            evidence_text="The adsorption energy of Li2S4 is -1.23 eV.",
            confidence=0.82,
        )
        claim = MechanismClaim(
            paper_id=paper.id,
            claim_type="adsorption",
            claim_text="The catalyst strengthens polysulfide adsorption.",
            evidence_types=["dft", "figure"],
            evidence_text="Charge redistribution indicates stronger adsorption.",
        )
        card = WritingCard(
            paper_id=paper.id,
            research_gap="Weak polysulfide adsorption remains a limitation.",
            proposed_solution="Use defect sites to tune adsorption.",
            core_hypothesis="Defect engineering improves sulfur conversion.",
        )
        session.add_all([section, figure, table, dft, claim, card])
        session.flush()
        session.add(
            EvidenceLocator(
                paper_id=paper.id,
                target_type="dft_result",
                target_id=str(dft.id),
                field_name="value",
                page=5,
                table_id=table.id,
                evidence_text=dft.evidence_text,
                locator_status="exact_page",
                locator_confidence=0.91,
                parser_source="test",
            )
        )
        session.commit()
        paper_id = str(paper.id)
        dft_id = str(dft.id)
        claim_id = str(claim.id)

    with mcp_auth_context(_ide_auth()):
        papers = query_papers(q="Ordinary IDE AI", limit=5)
        assert papers["returned"] == 1
        assert papers["items"][0]["id"] == paper_id

        context = get_codex_context(paper_id=paper_id)
        assert context["context"]["external_audit_precondition"]["status"] == "ready"
        assert len(context["context"]["content"]["sections"]) == 1
        assert len(context["context"]["content"]["figures"]) == 1
        assert len(context["context"]["content"]["tables"]) == 1

        dft_context = get_codex_item(paper_id=paper_id, item_type="dft_result", item_id=dft_id)
        assert dft_context["context"]["export_safety"]["eligible"] is False
        assert "missing_review" in dft_context["context"]["export_safety"]["blocked_reasons"]

        mechanism_context = get_codex_item(paper_id=paper_id, item_type="mechanism_claim", item_id=claim_id)
        assert mechanism_context["context"]["item"]["claim_text"].startswith("The catalyst strengthens")

        imported = import_analysis(
            paper_id=paper_id,
            source="assigned_dft_audit",
            source_label="Assigned AI DFT audit",
            raw_payload={
                "paper_id": paper_id,
                "agent_role": "dft_auditor",
                "verdict": "WARN",
                "recommended_action": "needs_human_review",
                "suspected_missing": [],
                "metadata_status": "ok",
                "section_structure_status": "ok",
                "table_status": "ok",
                "figure_status": "ok",
                "dft_status": "warn",
                "evidence_examples": [{"text": "DFT row needs final reviewer confirmation."}],
                "confidence": 0.74,
            },
        )
        assert imported["candidate_count"] == 1
        assert imported["candidates"][0]["type"] == "external_audit_opinion"

        with pytest.raises(PermissionError):
            approve_correction(str(UUID(int=0)))

    with Session(mcp_test_env["engine"]) as session:
        candidate = session.scalar(select(ExternalAnalysisCandidate))
        assert candidate is not None
        assert candidate.candidate_type == "external_audit_opinion"
        assert candidate.status == "candidate"
        assert candidate.materialized_target_type is None
        assert candidate.materialized_target_id is None
        assert candidate.normalized_payload["verification_status"] == "unverified"
        assert candidate.normalized_payload["source"] == "assigned_dft_audit"

        row = session.get(DFTResult, UUID(dft_id))
        assert row is not None
        assert row.candidate_status == "system_candidate"


def test_mcp_import_analysis_accepts_object_level_review_payload(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(title="MCP Object Audit Paper", pdf_path="mcp-object.pdf", authors=[])
        session.add(paper)
        session.flush()
        row = DFTResult(
            paper_id=paper.id,
            property_type="adsorption_energy",
            adsorbate="Li2S4",
            value=-1.2,
            unit="eV",
            evidence_text="Table 1 reports adsorption energy.",
            candidate_status="system_candidate",
        )
        session.add(row)
        session.commit()
        paper_id = str(paper.id)
        row_id = str(row.id)

    with mcp_auth_context(_auth()):
        imported = import_analysis(
            paper_id=paper_id,
            source="assigned_dft_audit",
            source_label="Assigned AI DFT audit",
            raw_payload={
                "object_review_audits": [
                    {
                        "target_type": "dft_results",
                        "target_id": row_id,
                        "field_name": "value",
                        "decision": "REVISE",
                        "evidence_checked": True,
                        "evidence_location": {"page": 8, "table": "Table 1"},
                        "corrected_value": -1.35,
                        "recommended_action": "propose_correction",
                        "confidence": 0.71,
                    }
                ]
            },
        )

    assert imported["candidate_count"] == 1
    candidate = imported["candidates"][0]
    assert candidate["type"] == "object_review_audit"
    assert candidate["target_type"] == "dft_results"
    assert candidate["target_id"] == row_id
    assert candidate["field_name"] == "value"
    assert candidate["decision"] == "REVISE"
    assert candidate["verification_status"] == "unverified"

    with Session(mcp_test_env["engine"]) as session:
        stored_row = session.get(DFTResult, UUID(row_id))
        stored_candidate = session.scalar(select(ExternalAnalysisCandidate))
        assert stored_row.candidate_status == "system_candidate"
        assert stored_candidate.candidate_type == "object_review_audit"
        assert stored_candidate.status == "candidate"
        assert stored_candidate.normalized_payload["writes_final_truth"] is False


def test_mcp_import_analysis_auto_applies_dual_ai_dft_reviews(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(title="MCP Auto Apply DFT Paper", pdf_path="mcp-auto-apply.pdf", authors=[])
        session.add(paper)
        session.flush()
        catalyst = CatalystSample(
            paper_id=paper.id,
            name="Vacancy graphene",
            catalyst_type="defective_graphene",
            coordination="single vacancy",
            support="graphene",
        )
        session.add(catalyst)
        session.flush()
        row = DFTResult(
            paper_id=paper.id,
            catalyst_sample_id=catalyst.id,
            property_type="adsorption_energy",
            adsorbate="Li2S4",
            value=-1.2,
            unit="eV",
            reaction_step="adsorption",
            evidence_text="Table 1 reports -1.20 eV for Li2S4 adsorption.",
            candidate_status="system_candidate",
        )
        session.add(row)
        session.flush()
        session.add(
            EvidenceLocator(
                paper_id=paper.id,
                target_type="dft_results",
                target_id=str(row.id),
                field_name="value",
                page=7,
                evidence_text="Table 1 reports -1.20 eV for Li2S4 adsorption.",
                locator_status="exact_page",
                locator_confidence=0.95,
                parser_source="test",
            )
        )
        session.commit()
        paper_id = str(paper.id)
        row_id = str(row.id)

    payload = {
        "object_review_audits": [
            {
                "target_type": "dft_results",
                "target_id": row_id,
                "field_name": "value",
                "decision": "PASS",
                "corrected_value": -1.2,
                "evidence_checked": True,
                "confidence": 0.91,
                "reason": "Table 1 confirms the DFT value.",
                "evidence_location": {"page": 7, "table": "Table 1", "quoted_text": "-1.20 eV"},
            }
        ]
    }

    with mcp_auth_context(_auth()):
        first = import_analysis(
            paper_id=paper_id,
            source="assigned_dft_audit",
            source_label="Assigned AI DFT audit A",
            raw_payload=payload,
        )
        second = import_analysis(
            paper_id=paper_id,
            source="assigned_dft_audit",
            source_label="Assigned AI DFT audit B",
            raw_payload=payload,
        )

    assert first["candidate_count"] == 1
    assert second["candidate_count"] == 1
    assert first["auto_apply_summary"]["object_reviews"]["pending_count"] == 1
    assert second["auto_apply_summary"]["object_reviews"]["applied_count"] == 1

    with Session(mcp_test_env["engine"]) as session:
        stored_row = session.get(DFTResult, UUID(row_id))
        candidates = session.query(ExternalAnalysisCandidate).order_by(ExternalAnalysisCandidate.created_at.asc()).all()
        reviews = session.query(ExtractionFieldReview).all()

    assert stored_row is not None
    assert stored_row.candidate_status == "ML_Ready"
    assert {candidate.status for candidate in candidates} == {"materialized"}
    assert reviews
    assert {review.reviewer_status for review in reviews} == {"verified"}


def test_mcp_get_dft_review_queue_returns_codex_ready_candidates(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(title="MCP DFT Queue Paper", doi="10.1000/dft-queue", year=2025, pdf_path="queue.pdf")
        session.add(paper)
        session.flush()
        row = DFTResult(
            paper_id=paper.id,
            property_type="reaction_barrier",
            adsorbate="vacancy",
            value=1.3,
            unit="eV",
            reaction_step="single vacancy migration",
            evidence_text="The migration barrier for the vacancy is 1.3 eV.",
            confidence=0.88,
        )
        session.add(row)
        session.flush()
        session.add(
            EvidenceLocator(
                paper_id=paper.id,
                target_type="dft_result",
                target_id=str(row.id),
                field_name="value",
                page=7,
                evidence_text=row.evidence_text,
                locator_status="exact_page",
                locator_confidence=0.93,
                parser_source="test",
            )
        )
        run = ExternalAnalysisRun(
            paper_id=paper.id,
            source="assigned_dft_audit",
            source_label="Assigned AI DFT audit",
            normalized_payload={"verdict": "WARN"},
            mapping_status="normalized",
        )
        session.add(run)
        session.flush()
        session.add(
            ExternalAnalysisCandidate(
                run_id=run.id,
                paper_id=paper.id,
                candidate_type="external_audit_opinion",
                normalized_payload={
                    "source": "assigned_dft_audit",
                    "source_label": "Assigned AI DFT audit",
                    "agent_role": "dft_auditor",
                    "model_name": "glm-test",
                    "verdict": "WARN",
                    "recommended_action": "verify_against_pdf",
                    "verification_status": "unverified",
                    "confidence": 0.72,
                    "summary": "Check the migration barrier against the source PDF.",
                },
                status="candidate",
                confidence=0.72,
            )
        )
        session.add(
            ExternalAnalysisCandidate(
                run_id=run.id,
                paper_id=paper.id,
                candidate_type="object_review_audit",
                normalized_payload={
                    "paper_id": str(paper.id),
                    "target_type": "dft_results",
                    "target_id": str(row.id),
                    "field_name": "value",
                    "source": "assigned_dft_audit",
                    "source_label": "Assigned AI DFT audit",
                    "agent_role": "dft_auditor",
                    "model_name": "glm-test",
                    "decision": "REVISE",
                    "recommended_action": "propose_correction",
                    "verification_status": "unverified",
                    "confidence": 0.71,
                    "reason": "Object-level check says the numeric value needs PDF review.",
                    "evidence_location": {"page": 7, "table": "Table 1"},
                    "writes_final_truth": False,
                    "human_confirmation_required": True,
                },
                status="candidate",
                confidence=0.71,
            )
        )
        session.commit()
        paper_id = str(paper.id)
        row_id = str(row.id)

    with mcp_auth_context(_auth()):
        payload = get_dft_review_queue(paper_id=paper_id, limit=10)

    assert payload["metadata"]["schema_version"] == "dft_review_queue_v1"
    assert payload["metadata"]["blocked_count"] == 1
    assert len(payload["rows"]) == 1
    row = payload["rows"][0]
    assert row["record_id"] == row_id
    assert row["blocked_reasons"] == ["missing_review"]
    assert row["recommended_action"] == "verify_against_pdf"
    assert row["sanity_flags"] == []
    assert row["can_mark_verified"] is True
    assert row["evidence_locators"][0]["page"] == 7
    assert row["primary_evidence_locator"]["page"] == 7
    assert row["evidence_page"] == 7
    assert row["pdf_page_url"].endswith(f"/api/papers/{paper_id}/pdf#page=7")
    assert row["latest_external_audit_opinions"][0]["source"] == "assigned_dft_audit"
    assert row["latest_external_audit_opinions"][0]["verification_status"] == "unverified"
    assert row["object_review_audits_count"] == 1
    assert row["object_review_audits"][0]["candidate_type"] == "object_review_audit"
    assert row["object_review_audits"][0]["decision"] == "REVISE"
    assert row["object_review_audits"][0]["verification_status"] == "unverified"
    assert row["object_review_audits"][0]["evidence_location"]["page"] == 7
    assert row["codex_item_url"].endswith(f"/codex-item/dft_result/{row_id}")
    assert row["correction_url"].endswith(f"/dft-results/{row_id}/corrections")

    with Session(mcp_test_env["engine"]) as session:
        stored_row = session.get(DFTResult, UUID(row_id))
        assert stored_row.candidate_status == "system_candidate"


def test_mcp_get_paper_knowledge_returns_section_fallback_candidates(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(
            title="MCP Knowledge Fallback Paper",
            abstract="Graphene vacancy defects alter adsorption and electronic structure.",
            pdf_path="knowledge.pdf",
        )
        session.add(paper)
        session.flush()
        session.add_all(
            [
                PaperNote(
                    paper_id=paper.id,
                    source="claude",
                    field_name="mechanism",
                    content="Check the vacancy adsorption mechanism before citing.",
                    quoted_text="vacancy defects alter adsorption",
                    page=1,
                ),
            ]
        )
        session.add(
            PaperSection(
                paper_id=paper.id,
                section_title="Results and Discussion",
                section_type="results",
                text="Vacancy defects alter adsorption energy and charge density around the defect site.",
                page_start=4,
                page_end=5,
            )
        )
        session.commit()
        paper_id = str(paper.id)

    with mcp_auth_context(_auth()):
        payload = get_paper_knowledge(paper_id=paper_id, max_candidates=10)

    assert payload["schema_version"] == "paper_knowledge_context_v1"
    assert payload["metadata"]["returned"] >= 2
    categories = {item["category"] for item in payload["candidates"]}
    assert "mechanism_context" in categories
    assert any(item["source_type"] == "paper_note" for item in payload["candidates"])
    assert payload["reliability_policy"]["knowledge_items_are_candidates"] is True


def test_mcp_insert_word_citation_creates_guarded_docx_copy(mcp_test_env):
    docx_path = mcp_test_env["tmpdir"] / "draft.docx"
    document = Document()
    document.add_paragraph("Draft manuscript body.")
    document.save(docx_path)

    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(
            title="MCP Word Citation Paper",
            year=2026,
            journal="Citation Journal",
            authors=[{"last": "Word"}],
            abstract="Graphene defect citation context.",
            pdf_path="word.pdf",
        )
        session.add(paper)
        session.commit()
        paper_id = str(paper.id)

    with mcp_auth_context(_auth()):
        payload = insert_word_citation(
            docx_path=str(docx_path),
            selected_paper_id=paper_id,
            text="Graphene defects alter adsorption behavior.",
            output_filename="mcp-word-citation.docx",
        )

    assert payload["status"] == "inserted"
    assert payload["output_filename"] == "mcp-word-citation.docx"
    assert payload["safety"]["mutates_original_file"] is False
    assert payload["safety"]["writes_database"] is False
    paragraphs = [paragraph.text for paragraph in Document(payload["output_path"]).paragraphs]
    assert paragraphs[0] == "Draft manuscript body."
    assert "[DRAFT CITATION - VERIFY SOURCE BEFORE USE: Word, 2026]" in paragraphs[-1]


def test_admin_mcp_reject_dft_result_leaves_active_queue(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(title="MCP Reject DFT Candidate", pdf_path="reject-dft.pdf")
        session.add(paper)
        session.flush()
        row = DFTResult(
            paper_id=paper.id,
            property_type="limiting_potential",
            adsorbate="[22]",
            value=436.0,
            unit="e",
            evidence_text="Reference-like artifact was parsed as a DFT result.",
        )
        session.add(row)
        session.flush()
        session.add(
            EvidenceLocator(
                paper_id=paper.id,
                target_type="dft_result",
                target_id=str(row.id),
                field_name="value",
                page=4,
                evidence_text=row.evidence_text,
                locator_status="exact_page",
                locator_confidence=0.9,
                parser_source="test",
            )
        )
        session.commit()
        paper_id = str(paper.id)
        row_id = str(row.id)

    with mcp_auth_context(_admin_auth()):
        rejected = reject_dft_result(
            paper_id=paper_id,
            dft_result_id=row_id,
            confirm_reject_candidate=True,
            reviewer_note="Reject citation-like DFT artifact.",
        )
        active_queue = get_dft_review_queue(paper_id=paper_id)
        rejected_queue = get_dft_review_queue(paper_id=paper_id, status="rejected")

    assert rejected["export_safety"]["review_status"] == "rejected"
    assert rejected["export_safety"]["blocked_reasons"] == ["unsafe_review"]
    assert active_queue["rows"] == []
    assert rejected_queue["rows"][0]["record_id"] == row_id
    assert rejected_queue["rows"][0]["decision_status"] == "rejected"


def test_mcp_propose_dft_result_correction_enters_review_queue(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(title="MCP DFT Correction Target", pdf_path="dft-correction.pdf")
        session.add(paper)
        session.flush()
        row = DFTResult(
            paper_id=paper.id,
            property_type="limiting_potential",
            adsorbate="ORR",
            value=0.66,
            unit="e",
            evidence_text="The limiting potential is 0.66 V.",
            confidence=0.81,
        )
        session.add(row)
        session.commit()
        paper_id = str(paper.id)
        row_id = str(row.id)

    with mcp_auth_context(_auth()):
        with pytest.raises(ValueError):
            propose_dft_result_correction(
                paper_id=paper_id,
                dft_result_id=row_id,
                field_name="unit",
                proposed_value="V",
                reason="The source table reports potential in volts.",
                confirm_correction_proposal=False,
            )
        correction = propose_dft_result_correction(
            paper_id=paper_id,
            dft_result_id=row_id,
            field_name="unit",
            proposed_value="V",
            reason="The source table reports potential in volts.",
            confirm_correction_proposal=True,
            evidence_payload={"page": 6, "quoted_text": "The limiting potential is 0.66 V."},
        )

    assert correction["status"] == "pending"
    assert correction["field_name"] == "dft_results"
    assert correction["target_path"] == f"dft_results:{row_id}:unit"
    assert correction["proposed_value"] == "V"

    with mcp_auth_context(_admin_auth()):
        approved = approve_correction(correction["id"])
        assert approved["status"] == "approved"

    with Session(mcp_test_env["engine"]) as session:
        updated = session.get(DFTResult, UUID(row_id))
        assert updated is not None
        assert updated.unit == "V"
        audit = session.scalar(select(AuditLog).where(AuditLog.action == "propose_dft_result_correction"))
        assert audit is not None
        assert audit.target_id == correction["id"]


def test_admin_mcp_review_flow_applies_or_rejects_corrections(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(title="Review Target", abstract="Old abstract", pdf_path="review.pdf")
        session.add(paper)
        session.commit()
        paper_id = str(paper.id)

    with mcp_auth_context(_auth()):
        first = propose_correction(
            paper_id=paper_id,
            field_name="abstract",
            target_path="abstract",
            operation="replace",
            proposed_value="Approved abstract",
            reason="Better aligned with PDF abstract.",
        )
        second = propose_correction(
            paper_id=paper_id,
            field_name="title",
            target_path="title",
            operation="replace",
            proposed_value="Rejected title",
            reason="This one should be rejected.",
        )

    with mcp_auth_context(_admin_auth()):
        queue = get_correction_queue()
        assert len(queue["items"]) == 2

        approved = approve_correction(first["id"])
        assert approved["status"] == "approved"
        assert approved["reviewed_by"] == "admin"

        rejected = reject_correction(second["id"], reason="Not supported by source PDF.")
        assert rejected["status"] == "rejected"
        assert rejected["reviewed_by"] == "admin"

    with Session(mcp_test_env["engine"]) as session:
        paper = session.get(Paper, UUID(first["paper_id"]))
        assert paper is not None
        assert paper.abstract == "Approved abstract"
        assert paper.title == "Review Target"

        logs = session.scalars(select(AuditLog).order_by(AuditLog.created_at.asc())).all()
        actions = [log.action for log in logs]
        assert "approve_correction" in actions
        assert "reject_correction" in actions


def test_admin_mcp_review_flow_applies_dft_result_patch(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(title="DFT Review Target", pdf_path="dft-review.pdf")
        session.add(paper)
        session.flush()
        dft_result = DFTResult(
            paper_id=paper.id,
            adsorbate="Li2S4",
            property_type="adsorption_energy",
            value=-1.23,
            unit="eV",
            confidence=0.82,
        )
        session.add(dft_result)
        session.commit()
        paper_id = str(paper.id)
        dft_result_id = str(dft_result.id)

    with mcp_auth_context(_auth()):
        correction = propose_correction(
            paper_id=paper_id,
            field_name="dft_results",
            target_path=f"dft_results:{dft_result_id}:value",
            operation="replace",
            proposed_value=-1.45,
            reason="Cross-check with Table 2 gives a corrected adsorption energy.",
            evidence_payload={
                "page": 6,
                "section_title": "DFT Results",
                "quoted_text": "Li2S4 adsorption energy on Fe-N4 is -1.45 eV.",
            },
        )
        assert correction["status"] == "pending"

    with mcp_auth_context(_admin_auth()):
        approved = approve_correction(correction["id"])
        assert approved["status"] == "approved"

    with Session(mcp_test_env["engine"]) as session:
        updated = session.get(DFTResult, UUID(dft_result_id))
        assert updated is not None
        assert updated.value == -1.45


def test_admin_mcp_review_flow_applies_mechanism_claim_patch(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(title="Mechanism Review Target", pdf_path="mechanism-review.pdf")
        session.add(paper)
        session.flush()
        claim = MechanismClaim(
            paper_id=paper.id,
            claim_type="shuttle_suppression",
            claim_text="Fe-N4 suppresses the shuttle effect.",
            evidence_types=["electrochem"],
            confidence=0.61,
        )
        session.add(claim)
        session.commit()
        paper_id = str(paper.id)
        claim_id = str(claim.id)

    with mcp_auth_context(_auth()):
        correction = propose_correction(
            paper_id=paper_id,
            field_name="mechanism_claims",
            target_path=f"mechanism_claims:{claim_id}:claim_text",
            operation="replace",
            proposed_value="Fe-N4 is associated with reduced shuttle behavior under the reported test conditions.",
            reason="The original wording overstates causality compared with the source text.",
        )

    with mcp_auth_context(_admin_auth()):
        detail = get_correction_detail(correction["id"])
        assert detail["current_value"] == "Fe-N4 suppresses the shuttle effect."
        assert detail["proposed_value"] == "Fe-N4 is associated with reduced shuttle behavior under the reported test conditions."

        approved = approve_correction(correction["id"])
        assert approved["status"] == "approved"

    with Session(mcp_test_env["engine"]) as session:
        updated = session.get(MechanismClaim, UUID(claim_id))
        assert updated is not None
        assert (
            updated.claim_text
            == "Fe-N4 is associated with reduced shuttle behavior under the reported test conditions."
        )


def test_http_correction_detail_returns_current_value_for_structured_targets(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(title="Detail Target", pdf_path="detail.pdf")
        session.add(paper)
        session.flush()
        perf = ElectrochemicalPerformance(
            paper_id=paper.id,
            capacity_value=873.0,
            cycle_number=100,
            rate="0.5 C",
            evidence_text="873 mAh g-1 at 100 cycles",
        )
        session.add(perf)
        session.flush()
        correction = PaperCorrection(
            paper_id=paper.id,
            source="claude",
            field_name="electrochemical_performance",
            target_path=f"electrochemical_performance:{perf.id}:capacity_value",
            operation="replace",
            proposed_value=892.0,
            reason="Figure annotation shows 892 mAh g-1 instead of 873.",
            status="pending",
        )
        session.add(correction)
        session.commit()
        correction_id = str(correction.id)

    client = TestClient(app)
    response = client.get(
        f"/api/corrections/{correction_id}",
        headers={"Authorization": "Bearer litmcp_admin"},
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload["current_value"] == 873.0
    assert payload["proposed_value"] == 892.0
    assert payload["target_exists"] is True


@pytest.mark.anyio
async def test_scan_local_pdfs_and_ingest_pdf_batch(mcp_test_env, monkeypatch):
    with tempfile.TemporaryDirectory() as tmpdir:
        folder = Path(tmpdir)
        first_pdf = folder / "paper_a.pdf"
        second_pdf = folder / "paper_b.pdf"
        first_pdf.write_bytes(b"%PDF-1.4 first")
        second_pdf.write_bytes(b"%PDF-1.4 second")

        with Session(mcp_test_env["engine"]) as session:
            existing = Paper(
                title="Already Parsed",
                pdf_path="stored_existing.pdf",
                source_path=str(first_pdf.resolve()),
            )
            session.add(existing)
            session.commit()
            existing_id = str(existing.id)

        async def fake_ingest_pdf(self, source_path, original_filename, copy_pdf=True, external_metadata=None, source_reference=None):
            paper = Paper(
                title=f"Parsed {original_filename}",
                pdf_path=f"stored_{original_filename}",
                source_path=source_reference,
            )
            self.session.add(paper)
            self.session.commit()
            self.session.refresh(paper)
            return paper

        monkeypatch.setattr("app.services.paper_ingestion.PaperIngestionService.ingest_pdf", fake_ingest_pdf)

        with mcp_auth_context(_auth()):
            scan = scan_local_pdfs(folder_path=str(folder), recursive=False, limit=10)
            assert scan["returned"] == 2
            existing_items = [item for item in scan["items"] if item["already_ingested"]]
            pending_items = [item for item in scan["items"] if not item["already_ingested"]]
            assert len(existing_items) == 1
            assert existing_items[0]["paper_id"] == existing_id
            assert len(pending_items) == 1
            assert pending_items[0]["filename"] == "paper_b.pdf"

            batch = await ingest_pdf_batch(
                folder_path=str(folder),
                recursive=False,
                limit=10,
                only_unparsed=True,
            )
            assert batch["requested"] == 2
            statuses = {item["path"]: item["status"] for item in batch["results"]}
            assert statuses[str(first_pdf.resolve())] == "already_ingested"
            assert statuses[str(second_pdf.resolve())] == "completed"

        with Session(mcp_test_env["engine"]) as session:
            rows = session.scalars(select(Paper).order_by(Paper.created_at.asc())).all()
            assert len(rows) == 2
            assert any(row.source_path == str(second_pdf.resolve()) for row in rows)


@pytest.mark.anyio
async def test_parse_paper_reuses_existing_paper_and_records_job(mcp_test_env, monkeypatch):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(
            doi="10.1000/existing-doi",
            title="Existing Paper",
            pdf_path="existing.pdf",
        )
        session.add(paper)
        session.commit()
        existing_paper_id = str(paper.id)

    class FakeDiscoveryService:
        def fetch_metadata(self, identifier, providers=None):
            return object(), {
                "doi": "10.1000/existing-doi",
                "title": "Existing Paper",
                "authors": [],
                "providers": providers or [],
            }

    monkeypatch.setattr("app.mcp.server.DiscoveryService", FakeDiscoveryService)

    with mcp_auth_context(_auth()):
        job = await parse_paper(identifier="10.1000/existing-doi", providers=["openalex"])
        assert job["status"] == "completed"
        assert job["paper_id"] == existing_paper_id

        fetched = get_parse_status(job_id=job["id"])
        assert fetched["identifier"] == "10.1000/existing-doi"
        assert fetched["status"] == "completed"


def test_mcp_http_auth_middleware_requires_api_key(mcp_test_env):
    with TestClient(app) as client:
        response = client.get("/mcp")
        assert response.status_code == 401
        assert response.json()["detail"] == "Missing MCP API key"

        authorized = client.get("/mcp", headers={"Authorization": "Bearer litmcp_claude"})
        assert authorized.status_code != 401


def test_http_correction_review_api_requires_admin_and_applies_update(mcp_test_env):
    with Session(mcp_test_env["engine"]) as session:
        paper = Paper(title="HTTP Review Target", abstract="Initial abstract", pdf_path="http-review.pdf")
        session.add(paper)
        session.flush()
        correction = PaperCorrection(
            paper_id=paper.id,
            source="claude",
            field_name="abstract",
            target_path="abstract",
            operation="replace",
            proposed_value="HTTP approved abstract",
            reason="Reviewed against source abstract.",
            status="pending",
        )
        session.add(correction)
        session.commit()
        correction_id = str(correction.id)
        paper_id = str(paper.id)

    client = TestClient(app)
    forbidden = client.post(
        f"/api/corrections/{correction_id}/approve",
        headers={"Authorization": "Bearer litmcp_claude"},
    )
    assert forbidden.status_code == 403

    approved = client.post(
        f"/api/corrections/{correction_id}/approve",
        headers={"Authorization": "Bearer litmcp_admin"},
    )
    assert approved.status_code == 200
    assert approved.json()["status"] == "approved"

    queue = client.get(
        "/api/corrections",
        headers={"Authorization": "Bearer litmcp_admin"},
    )
    assert queue.status_code == 200
    assert queue.json() == []

    with Session(mcp_test_env["engine"]) as session:
        paper = session.get(Paper, UUID(paper_id))
        assert paper is not None
        assert paper.abstract == "HTTP approved abstract"
