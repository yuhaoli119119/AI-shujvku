from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from sqlalchemy.orm import Session

from app.config import Settings
from app.services.writing_citation_insertion_service import (
    CitationInsertionDraftRequest,
    WritingCitationInsertionService,
)


@dataclass(frozen=True)
class WordCitationInsertRequest:
    document_bytes: bytes
    filename: str
    text: str
    selected_paper_id: UUID
    citation_marker: str | None = None
    docx_insertion_mode: str = "append_paragraph"
    citation_insertion_mode: str = "parenthetical"
    citation_style: str = "draft_author_year"
    placeholder: str | None = None
    output_filename: str | None = None
    user_note: str | None = None


class WordCitationInsertionService:
    """Insert a citation draft into a Word document copy without mutating the DB."""

    allowed_docx_modes = {"append_paragraph", "replace_placeholder"}
    allowed_citation_modes = {"parenthetical", "narrative", "comment_only"}
    allowed_citation_styles = {"draft_author_year", "placeholder"}

    def __init__(self, session: Session, settings: Settings) -> None:
        self.session = session
        self.settings = settings

    def insert(self, request: WordCitationInsertRequest) -> dict[str, Any] | None:
        self._validate_request(request)
        draft = WritingCitationInsertionService(self.session).draft(
            CitationInsertionDraftRequest(
                text=request.text,
                selected_paper_id=request.selected_paper_id,
                citation_marker=request.citation_marker,
                insertion_mode=request.citation_insertion_mode,
                citation_style=request.citation_style,
                user_note=request.user_note,
            )
        )
        if draft is None:
            return None
        if draft.get("draft_text") is None:
            return self._blocked_response(draft, request)

        try:
            document = Document(BytesIO(request.document_bytes))
        except Exception as exc:  # pragma: no cover - docx internals vary by version
            raise ValueError(f"Unable to read DOCX file: {exc}") from exc

        inserted_text = str(draft["draft_text"])
        replaced_count = 0
        placeholder = self._placeholder(request)
        if request.docx_insertion_mode == "append_paragraph":
            document.add_paragraph(inserted_text)
        elif request.docx_insertion_mode == "replace_placeholder":
            replacement = str(draft["citation_marker"])
            replaced_count = self._replace_placeholder(document, placeholder, replacement)
            if replaced_count <= 0:
                raise ValueError(f"Placeholder not found in DOCX: {placeholder}")
            inserted_text = replacement
        else:
            raise ValueError("Unsupported docx_insertion_mode")

        output_path = self._output_path(request.filename, request.output_filename)
        document.save(output_path)
        relative_path = self._relative_output_path(output_path)
        return {
            "status": "inserted",
            "paper_id": draft.get("paper_id"),
            "title": draft.get("title"),
            "output_path": str(output_path.resolve()),
            "output_relative_path": relative_path,
            "output_filename": output_path.name,
            "download_url": f"/api/writing/word/exports/{output_path.name}",
            "docx_insertion_mode": request.docx_insertion_mode,
            "citation_insertion_mode": request.citation_insertion_mode,
            "placeholder": placeholder if request.docx_insertion_mode == "replace_placeholder" else None,
            "placeholder_replaced_count": replaced_count,
            "inserted_text": inserted_text,
            "draft": draft,
            "warnings": draft.get("warnings", []),
            "safety": self._safety(draft),
        }

    def _validate_request(self, request: WordCitationInsertRequest) -> None:
        if not request.filename.lower().endswith(".docx"):
            raise ValueError("Only .docx files are supported")
        if not request.document_bytes:
            raise ValueError("DOCX file is empty")
        if not request.text.strip():
            raise ValueError("text must not be blank")
        if request.docx_insertion_mode not in self.allowed_docx_modes:
            raise ValueError("docx_insertion_mode must be append_paragraph or replace_placeholder")
        if request.citation_insertion_mode not in self.allowed_citation_modes:
            raise ValueError("citation_insertion_mode must be parenthetical, narrative, or comment_only")
        if request.citation_style not in self.allowed_citation_styles:
            raise ValueError("citation_style must be draft_author_year or placeholder")

    def _blocked_response(self, draft: dict[str, Any], request: WordCitationInsertRequest) -> dict[str, Any]:
        return {
            "status": "blocked",
            "paper_id": draft.get("paper_id"),
            "title": draft.get("title"),
            "output_path": None,
            "output_relative_path": None,
            "output_filename": None,
            "download_url": None,
            "docx_insertion_mode": request.docx_insertion_mode,
            "citation_insertion_mode": request.citation_insertion_mode,
            "placeholder": self._placeholder(request) if request.docx_insertion_mode == "replace_placeholder" else None,
            "placeholder_replaced_count": 0,
            "inserted_text": None,
            "draft": draft,
            "warnings": draft.get("warnings", []),
            "safety": self._safety(draft),
        }

    def _output_path(self, input_filename: str, output_filename: str | None) -> Path:
        export_dir = self.settings.storage_root / "word_exports"
        export_dir.mkdir(parents=True, exist_ok=True)
        name = output_filename.strip() if output_filename and output_filename.strip() else ""
        if not name:
            stem = Path(input_filename).stem or "document"
            stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
            name = f"{stem}_cited_{stamp}.docx"
        if not name.lower().endswith(".docx"):
            name = f"{name}.docx"
        safe_name = self._safe_filename(name)
        path = (export_dir / safe_name).resolve()
        export_root = export_dir.resolve()
        if not str(path).lower().startswith(str(export_root).lower()):
            raise ValueError("Invalid output filename")
        if path.exists():
            path = export_root / f"{path.stem}_{datetime.now(timezone.utc).strftime('%H%M%S%f')}.docx"
        return path

    def _relative_output_path(self, output_path: Path) -> str:
        try:
            return output_path.resolve().relative_to(self.settings.storage_root.resolve()).as_posix()
        except ValueError:
            return output_path.name

    @classmethod
    def resolve_export_path(cls, settings: Settings, filename: str) -> Path:
        safe_name = cls._safe_filename(filename)
        export_root = (settings.storage_root / "word_exports").resolve()
        path = (export_root / safe_name).resolve()
        if not str(path).lower().startswith(str(export_root).lower()):
            raise ValueError("Invalid export filename")
        if not path.exists() or not path.is_file():
            raise FileNotFoundError(safe_name)
        return path

    @staticmethod
    def _safe_filename(value: str) -> str:
        cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", Path(value).name).strip("._")
        return cleaned or "document_cited.docx"

    @staticmethod
    def _placeholder(request: WordCitationInsertRequest) -> str:
        if request.placeholder and request.placeholder.strip():
            return request.placeholder.strip()
        return f"{{{{CITE:{request.selected_paper_id}}}}}"

    def _replace_placeholder(self, document: Any, placeholder: str, replacement: str) -> int:
        count = 0
        for paragraph in self._iter_paragraphs(document):
            count += self._replace_in_paragraph(paragraph, placeholder, replacement)
        return count

    @staticmethod
    def _iter_paragraphs(document: Any):
        yield from document.paragraphs
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
        for section in document.sections:
            yield from section.header.paragraphs
            yield from section.footer.paragraphs

    @staticmethod
    def _replace_in_paragraph(paragraph: Any, placeholder: str, replacement: str) -> int:
        replaced = 0
        for run in paragraph.runs:
            if placeholder in run.text:
                replaced += run.text.count(placeholder)
                run.text = run.text.replace(placeholder, replacement)
        if replaced:
            return replaced
        if placeholder not in paragraph.text:
            return 0

        updated = paragraph.text.replace(placeholder, replacement)
        replaced = paragraph.text.count(placeholder)
        if paragraph.runs:
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(updated)
        return replaced

    @staticmethod
    def _safety(draft: dict[str, Any]) -> dict[str, Any]:
        return {
            "generated_docx_copy": True,
            "mutates_original_file": False,
            "writes_database": False,
            "marks_verified": False,
            "generates_bibliography": False,
            "requires_human_verification": bool(draft.get("requires_human_verification")),
            "can_insert_as_confirmed_citation": bool(draft.get("can_insert_as_confirmed_citation")),
            "final_citation_only_when_safe_verified": True,
        }
